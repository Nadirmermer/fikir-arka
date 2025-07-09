"""
AI Content API Endpoints
Generate and manage AI content
"""

from fastapi import APIRouter, HTTPException, BackgroundTasks, Response
from fastapi.responses import FileResponse
from typing import List
from sqlalchemy import select
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import tempfile

from app.database import get_db
from app.models import AIContent, AIContentCreate, AIContentResponse, Topic
from app.services.ai_service import AIService
import logging

router = APIRouter()
logger = logging.getLogger(__name__)
PROMPT_FILE_PATH = "ai_prompts/master_prompt.txt"

@router.get("/", response_model=List[AIContentResponse])
async def get_ai_contents():
    """Get all AI generated contents"""
    async with get_db() as db:
        result = await db.execute(select(AIContent).order_by(AIContent.created_at.desc()))
        contents = result.scalars().all()
        return contents

@router.post("/generate")
async def generate_ai_content(content_data: AIContentCreate, background_tasks: BackgroundTasks):
    """Generate AI content from liked topic"""
    # Verify topic exists and is liked
    async with get_db() as db:
        result = await db.execute(select(Topic).where(Topic.id == content_data.topic_id))
        topic = result.scalar_one_or_none()
        
        if not topic:
            raise HTTPException(status_code=404, detail="Topic not found")
        
        if topic.status != "liked":
            raise HTTPException(status_code=400, detail="Topic must be liked to generate content")
    
    # Create AI content record
    ai_content = AIContent(
        title=content_data.title,
        content=content_data.content,
        topic_id=content_data.topic_id,
        ai_model=content_data.ai_model,
        temperature=content_data.temperature,
        prompt_used=content_data.prompt_used,
        status="generating"
    )
    
    async with get_db() as db:
        db.add(ai_content)
        await db.commit()
        await db.refresh(ai_content)
    
    # Start generation in background
    background_tasks.add_task(_generate_ai_content, ai_content.id, topic.title, topic.description or topic.content or "")
    
    return {"success": True, "ai_content_id": ai_content.id, "message": "AI generation started"}

async def _generate_ai_content(ai_content_id: str, topic_title: str, topic_content: str):
    """Background task for AI content generation"""
    ai_service = AIService()
    
    async with get_db() as db:
        result = await db.execute(select(AIContent).where(AIContent.id == ai_content_id))
        ai_content = result.scalar_one_or_none()
        
        if not ai_content:
            return
        
        # Generate content
        generation_result = await ai_service.generate_content(topic_title, topic_content)
        
        if generation_result["success"]:
            ai_content.generated_content = generation_result["generated_content"]
            ai_content.status = "completed"
            ai_content.generation_time_seconds = generation_result["metadata"]["generation_time_seconds"]
            ai_content.content_length = len(generation_result["generated_content"])
        else:
            ai_content.status = "failed"
            ai_content.generated_content = f"Error: {generation_result['error']}"
        
        db.add(ai_content)
        await db.commit()

@router.get("/{ai_content_id}/export/word")
async def export_ai_content_to_word(ai_content_id: str):
    """Export AI content to Word document"""
    async with get_db() as db:
        result = await db.execute(select(AIContent).where(AIContent.id == ai_content_id))
        ai_content = result.scalar_one_or_none()
        
        if not ai_content:
            raise HTTPException(status_code=404, detail="AI content not found")
        
        if ai_content.status != "completed":
            raise HTTPException(status_code=400, detail="AI content is not completed yet")
    
    try:
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(ai_content.title, 0)
        title.alignment = 1  # Center alignment
        
        # Add metadata
        doc.add_heading('İçerik Bilgileri', level=1)
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Table Grid'
        
        metadata_table.cell(0, 0).text = 'Oluşturma Tarihi:'
        metadata_table.cell(0, 1).text = ai_content.created_at.strftime('%d/%m/%Y %H:%M')
        
        metadata_table.cell(1, 0).text = 'AI Model:'
        metadata_table.cell(1, 1).text = ai_content.ai_model
        
        metadata_table.cell(2, 0).text = 'İçerik Uzunluğu:'
        metadata_table.cell(2, 1).text = f'{ai_content.content_length} karakter'
        
        metadata_table.cell(3, 0).text = 'Üretim Süresi:'
        metadata_table.cell(3, 1).text = f'{ai_content.generation_time_seconds} saniye'
        
        # Add generated content
        doc.add_heading('YouTube Video Scripti', level=1)
        content_para = doc.add_paragraph(ai_content.generated_content)
        content_para.alignment = 0  # Left alignment
        
        # Add original topic info
        doc.add_heading('Kaynak Bilgisi', level=1)
        source_para = doc.add_paragraph(f"Orijinal İçerik: {ai_content.content}")
        source_para.alignment = 0
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            tmp_filename = tmp_file.name
        
        # Create safe filename
        safe_title = "".join(c for c in ai_content.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        filename = f"{safe_title[:50]}_AI_Content_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        
        return FileResponse(
            path=tmp_filename,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            background=lambda: os.unlink(tmp_filename)  # Clean up temp file after sending
        )
        
    except Exception as e:
        logger.error(f"Word export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Word export failed: {str(e)}")

@router.delete("/{ai_content_id}")
async def delete_ai_content(ai_content_id: str):
    """Delete AI content"""
    async with get_db() as db:
        result = await db.execute(select(AIContent).where(AIContent.id == ai_content_id))
        ai_content = result.scalar_one_or_none()
        
        if not ai_content:
            raise HTTPException(status_code=404, detail="AI content not found")
        
        await db.delete(ai_content)
        await db.commit()
        
        return {"success": True, "message": "AI content deleted successfully"}

@router.get("/test")
async def test_ai_connection():
    """Test AI service connection"""
    ai_service = AIService()
    return await ai_service.test_connection() 

@router.get("/prompt")
async def get_ai_prompt():
    """Get the current master AI prompt"""
    try:
        with open(PROMPT_FILE_PATH, 'r', encoding='utf-8') as f:
            prompt = f.read()
        return {"prompt": prompt}
    except FileNotFoundError:
        return {"prompt": ""}
    except Exception as e:
        logger.error(f"Error reading prompt file: {e}")
        raise HTTPException(status_code=500, detail="Could not read prompt file")

@router.post("/prompt")
async def save_ai_prompt(request: dict):
    """Save the master AI prompt"""
    prompt = request.get('prompt')
    if prompt is None:
        raise HTTPException(status_code=400, detail="Prompt content is missing")
    try:
        # Ensure directory exists
        os.makedirs("ai_prompts", exist_ok=True)
        with open(PROMPT_FILE_PATH, 'w', encoding='utf-8') as f:
            f.write(prompt)
        return {"success": True, "message": "Prompt saved successfully"}
    except Exception as e:
        logger.error(f"Error writing prompt file: {e}")
        raise HTTPException(status_code=500, detail="Could not save prompt file") 